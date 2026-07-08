#!/usr/bin/env python3
"""为中文课题申报书 docx 套用公文格式（字体/字号/行距/对齐/缩进/颜色）。

默认值（三号=16pt，可用 CLI 参数覆盖，无需改脚本）：
  一级标题/文档标题 -> 黑体、加粗、黑色
  二级标题          -> 楷体_GB2312、加粗、黑色
  三级及以下标题     -> 楷体_GB2312、黑色
  正文              -> 仿宋_GB2312、黑色、行距 28 磅(EXACTLY)、两端对齐、
                      首行缩进 2 字(=2×字号)；列表段不加首行缩进

依赖：pip install python-docx
用法：
  python format_docx.py in.docx                 # 输出 in.formatted.docx（不覆盖原件）
  python format_docx.py in.docx -o out.docx     # 指定输出
  python format_docx.py in.docx --in-place       # 原地覆盖
  python format_docx.py in.docx --h1-font 黑体 --body-font 仿宋 \
      --size 16 --line-spacing 28 --indent 32    # 按本课题格式要求调参
"""
import argparse
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor


def set_run_font(run, cn_font, size, bold=False):
    """把中英文字体都设为 cn_font，并设字号/颜色/加粗。"""
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), cn_font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.bold = bold


def is_list_paragraph(p):
    """列表段判定：直接带 numPr，或样式名指示列表（List.../列表/List Paragraph）。

    pandoc 生成的列表通常带直接 numPr；用命名列表样式的文档靠样式名兜底。
    """
    ppr = p._p.pPr
    if ppr is not None and ppr.find(qn("w:numPr")) is not None:
        return True
    name = (p.style.name or "") if p.style else ""
    return name.startswith("List") or "列表" in name


def heading_level(style_name):
    """返回标题级别：Title/标题->0，Heading N/标题 N->N，正文->None。

    兼容 pandoc 生成的英文样式名与中文模板本地化样式名（"标题 1"等）。
    """
    if not style_name:
        return None
    if style_name in ("Title", "标题"):
        return 0
    m = re.match(r"(?:Heading|标题)\s*(\d+)", style_name)
    if m:
        return int(m.group(1))
    return None


def format_paragraph(p, cfg):
    level = heading_level(p.style.name if p.style else None)
    if level in (0, 1):                       # 文档标题 / 一级标题
        font, bold = cfg.h1_font, True
    elif level == 2:                          # 二级标题
        font, bold = cfg.h2_font, True
    elif level is not None:                   # 三级及以下标题
        font, bold = cfg.h2_font, False
    else:                                     # 正文
        font, bold = cfg.body_font, False
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(cfg.line_spacing)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if not is_list_paragraph(p):
            pf.first_line_indent = Pt(cfg.indent)

    for run in p.runs:
        set_run_font(run, font, cfg.size, bold=bold)


def main():
    ap = argparse.ArgumentParser(description="申报书 docx 公文格式化")
    ap.add_argument("path")
    ap.add_argument("-o", "--output", help="输出文件路径")
    ap.add_argument("--in-place", action="store_true",
                    help="原地覆盖原文件（默认写到 *.formatted.docx）")
    ap.add_argument("--h1-font", dest="h1_font", default="黑体")
    ap.add_argument("--h2-font", dest="h2_font", default="楷体_GB2312",
                    help="二三级标题字体")
    ap.add_argument("--body-font", dest="body_font", default="仿宋_GB2312")
    ap.add_argument("--size", type=float, default=16, help="字号 pt，三号=16")
    ap.add_argument("--line-spacing", dest="line_spacing", type=float,
                    default=28, help="正文行距 pt（固定值）")
    ap.add_argument("--indent", type=float, default=32,
                    help="正文首行缩进 pt（2 字≈2×字号）")
    cfg = ap.parse_args()

    doc = Document(cfg.path)
    for p in doc.paragraphs:
        format_paragraph(p, cfg)
    # 表格内文字统一为正文字体（如仍保留少量表格）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        set_run_font(run, cfg.body_font, cfg.size)

    if cfg.output:
        out = cfg.output
    elif cfg.in_place:
        out = cfg.path
    else:
        base, ext = os.path.splitext(cfg.path)
        out = base + ".formatted" + ext
    doc.save(out)
    print(f"已套用公文格式 -> {out}")


if __name__ == "__main__":
    main()
